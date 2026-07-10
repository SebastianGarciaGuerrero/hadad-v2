"""
Documentos Word (.docx) formales para entregar a clínicas, deudores o
tribunales. Diseño sobrio de estudio jurídico (membrete Hadad & Asociados).

  GET /api/documentos/informe-gestiones/{cobranza_id} → Word con el historial
      completo de gestiones del caso a la fecha.
  GET /api/documentos/estado-cuenta/{cobranza_id}     → Word con el estado de
      cuenta actualizado (deuda, abonos, acuerdo, cuotas).
"""

from io import BytesIO
from uuid import UUID
from datetime import date, datetime

from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from app.database import get_db
from app.security import get_current_user
from app.models.cobranza import Cobranza
from app.models.gestion import Gestion, TipoGestion
from app.models.acuerdo import AcuerdoPago
from app.models.pago import Pago


router = APIRouter(
    prefix="/api/documentos",
    tags=["Documentos Word"],
    dependencies=[Depends(get_current_user)],
)

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

GRIS_OSCURO = RGBColor(0x21, 0x21, 0x26)
GRIS_SUAVE = RGBColor(0x71, 0x71, 0x7A)

MESES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio", "julio",
    "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]


def _clp(valor) -> str:
    return "$" + f"{int(valor):,}".replace(",", ".")


def _fecha_larga(f: date) -> str:
    return f"{f.day} de {MESES[f.month - 1]} de {f.year}"


def _fecha_corta(f) -> str:
    return f.strftime("%d-%m-%Y") if f else "—"


def _membrete(doc: Document):
    """Encabezado institucional: wordmark centrado + línea divisoria."""
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titulo.add_run("HADAD & ASOCIADOS")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = GRIS_OSCURO

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("ASESORÍA LEGAL Y FINANCIERA")
    r.font.size = Pt(9)
    r.font.color.rgb = GRIS_SUAVE
    # espaciado entre letras (estilo del logo)
    r.font.name = "Segoe UI"

    linea = doc.add_paragraph()
    linea.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = linea.add_run("_" * 72)
    r.font.color.rgb = GRIS_SUAVE
    r.font.size = Pt(8)


def _titulo_documento(doc: Document, texto: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = GRIS_OSCURO


def _tabla_datos(doc: Document, pares: list):
    """Tabla de dos columnas etiqueta/valor con estilo sobrio."""
    tabla = doc.add_table(rows=len(pares), cols=2)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (etiqueta, valor) in enumerate(pares):
        c0, c1 = tabla.rows[i].cells
        c0.width, c1.width = Cm(5.5), Cm(10.5)
        r = c0.paragraphs[0].add_run(etiqueta)
        r.bold = True
        r.font.size = Pt(10)
        r = c1.paragraphs[0].add_run(str(valor))
        r.font.size = Pt(10)
    return tabla


def _pie_firma(doc: Document):
    doc.add_paragraph()
    doc.add_paragraph()
    firma = doc.add_paragraph()
    firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = firma.add_run("_______________________________")
    r.font.color.rgb = GRIS_OSCURO
    quien = doc.add_paragraph()
    quien.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = quien.add_run("González & Hadad Profesionales Asociados")
    r.bold = True
    r.font.size = Pt(10)
    dire = doc.add_paragraph()
    dire.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = dire.add_run("Valparaíso, Chile")
    r.font.size = Pt(9)
    r.font.color.rgb = GRIS_SUAVE


def _cargar_cobranza(db: Session, cobranza_id: UUID) -> Cobranza:
    cobranza = db.query(Cobranza).filter(Cobranza.id == cobranza_id).first()
    if not cobranza:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Cobranza con id {cobranza_id} no encontrada",
        )
    return cobranza


def _responder_docx(doc: Document, nombre: str) -> StreamingResponse:
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type=DOCX_MIME,
        headers={"Content-Disposition": f'attachment; filename="{nombre}"'},
    )


@router.get("/informe-gestiones/{cobranza_id}")
def informe_gestiones(cobranza_id: UUID, db: Session = Depends(get_db)):
    """Word formal con todas las gestiones del caso a la fecha."""
    cob = _cargar_cobranza(db, cobranza_id)
    gestiones = (
        db.query(Gestion)
        .filter(Gestion.cobranza_id == cobranza_id)
        .order_by(Gestion.fecha_gestion)
        .all()
    )
    tipos = {t.id: t.nombre for t in db.query(TipoGestion).all()}

    doc = Document()
    _membrete(doc)
    _titulo_documento(doc, "INFORME DE GESTIONES DE COBRANZA")

    intro = doc.add_paragraph()
    r = intro.add_run(
        f"En Valparaíso, a {_fecha_larga(date.today())}, se informa el detalle "
        f"de las gestiones de cobranza realizadas a la fecha respecto del "
        f"siguiente caso:"
    )
    r.font.size = Pt(10.5)

    _tabla_datos(doc, [
        ("N° de cobranza", cob.numero),
        ("ID cliente", cob.id_clinica or "—"),
        ("Deudor", f"{cob.deudor.nombre} — RUT {cob.deudor.rut}" if cob.deudor else "—"),
        ("Cliente", (cob.cliente.nombre_fantasia or cob.cliente.razon_social) if cob.cliente else "—"),
        ("Filial", cob.filial.nombre if cob.filial else "—"),
        ("Deuda original", _clp(cob.monto_original)),
        ("Saldo a la fecha", _clp(cob.monto_actual)),
    ])

    doc.add_paragraph()
    sub = doc.add_paragraph()
    r = sub.add_run(f"Detalle de gestiones ({len(gestiones)})")
    r.bold = True
    r.font.size = Pt(11)

    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = "Table Grid"
    encabezados = ("Fecha", "Tipo", "Detalle de la gestión")
    anchos = (Cm(3), Cm(3.6), Cm(9.4))
    for i, (texto, ancho) in enumerate(zip(encabezados, anchos)):
        celda = tabla.rows[0].cells[i]
        celda.width = ancho
        r = celda.paragraphs[0].add_run(texto)
        r.bold = True
        r.font.size = Pt(10)
    for g in gestiones:
        fila = tabla.add_row().cells
        for i, ancho in enumerate(anchos):
            fila[i].width = ancho
        fila[0].paragraphs[0].add_run(
            g.fecha_gestion.strftime("%d-%m-%Y")).font.size = Pt(9.5)
        fila[1].paragraphs[0].add_run(
            tipos.get(g.tipo_id, "Gestión")).font.size = Pt(9.5)
        fila[2].paragraphs[0].add_run(g.descripcion).font.size = Pt(9.5)

    _pie_firma(doc)

    nombre = f"informe_gestiones_cobranza_{cob.numero}_{date.today().isoformat()}.docx"
    return _responder_docx(doc, nombre)


# ------------------------------------------------------------
# Estado de cuenta: carta formal calcada de la "BASE CARTAS
# REDSALUD" real del estudio (Courier New, estilo legal).
# ------------------------------------------------------------

ABREV_DOCUMENTO = {
    "pagare": "PA", "factura": "FA", "letra": "LE", "cheque": "CH", "otro": "OT",
}

PIE_CARTA = (
    "Errázuriz N° 1178, oficina 74, Valparaíso (Atención de 10 a 16 hrs.). "
    "Fono: (32)2450990 – (9)79593717\n"
    " gisellerojas.hadadyasociados@gmail.com; recepcion@hadadyasociados.cl "
    "www.hadadyasociados.cl"
)


def _rut_con_puntos(rut) -> str:
    """'5743070-2' → '5.743.070-2' (formato de la carta)."""
    if not rut:
        return "—"
    cuerpo, _, dv = str(rut).partition("-")
    cuerpo = cuerpo.replace(".", "")
    con_puntos = f"{int(cuerpo):,}".replace(",", ".") if cuerpo.isdigit() else cuerpo
    return f"{con_puntos}-{dv}" if dv else con_puntos


def _miles(valor) -> str:
    """Monto con separador de miles y sin símbolo (como la carta)."""
    return f"{int(valor):,}".replace(",", ".")


def _linea(doc: Document, texto: str, tam: float = 8.5, negrita: bool = False,
           alineacion=None, justificar: bool = False):
    """Párrafo en Courier New (la carta usa fuente monoespaciada)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    if alineacion is not None:
        p.alignment = alineacion
    if justificar:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(texto)
    r.font.name = "Courier New"
    r.font.size = Pt(tam)
    r.bold = negrita
    return p


@router.get("/estado-cuenta/{cobranza_id}")
def estado_cuenta(cobranza_id: UUID, db: Session = Depends(get_db)):
    """
    Carta de estado de cuenta con los datos actualizados del caso,
    con el mismo formato de la carta que el estudio envía (BASE REDSALUD).
    """
    from docx.shared import Cm as _Cm
    from app.models.paciente import Paciente

    cob = _cargar_cobranza(db, cobranza_id)
    pagos = db.query(Pago).filter(Pago.cobranza_id == cobranza_id).all()

    deudor = cob.deudor
    cliente = cob.cliente
    paciente = None
    if cob.paciente_id:
        paciente = db.query(Paciente).filter(Paciente.id == cob.paciente_id).first()

    # Celular del deudor (primer contacto activo tipo celular/teléfono).
    celular = ""
    if deudor is not None:
        for c in deudor.contactos:
            if c.activo and c.tipo in ("celular", "telefono", "whatsapp"):
                celular = c.valor
                break

    # Cifras del cuadro resumen. El SALDO suma lo que resta de capital más
    # honorarios/intereses/gastos informados de la cobranza.
    nominal = int(cob.monto_original)
    gastos = int(cob.gastos_hadad or 0)
    abonos = int(sum((p.capital_clinica or 0) for p in pagos))
    honorarios = int(cob.honorarios_hadad or 0)
    intereses = int(cob.intereses_hadad or 0)
    saldo = int(cob.monto_actual) + honorarios + intereses + gastos

    doc = Document()
    seccion = doc.sections[0]
    seccion.top_margin = _Cm(2.5)
    seccion.left_margin = _Cm(3.0)
    seccion.right_margin = _Cm(2.0)

    hoy = date.today()
    _linea(doc, f"{'':>49}Valparaíso, {hoy.day} de {MESES[hoy.month - 1]}, {hoy.year}")
    _linea(doc, "")
    _linea(doc, "")
    _linea(doc, "          Señor(es)")
    id_texto = f" ID {cob.id_clinica}" if cob.id_clinica else ""
    nombre_deudor = (deudor.nombre.upper() if deudor else "—") + id_texto
    _linea(doc, f"         {nombre_deudor:<52}N° COB.: {cob.numero}")
    _linea(doc, f"         R.U.T.: {_rut_con_puntos(deudor.rut if deudor else '')}")
    if paciente is not None:
        _linea(doc, f"         PACIENTE {paciente.nombre.upper()}")
    if celular:
        _linea(doc, f"{'':>59}N° CEL.: {celular}")
    if deudor is not None and deudor.direccion:
        _linea(doc, f"         {deudor.direccion.upper()}")
    ciudad = (deudor.ciudad or deudor.comuna) if deudor else None
    if ciudad:
        _linea(doc, f"         {ciudad.upper()}")
    _linea(doc, "")
    _linea(doc, "         Estimado Señor(es):")
    _linea(doc, "")

    estudio = (cliente.nombre_fantasia or cliente.razon_social).upper() if cliente else "—"
    razon = cliente.razon_social.upper() if cliente else "—"
    rut_cliente = _rut_con_puntos(cliente.rut) if cliente else "—"
    cuerpo = (
        f"         Nuestro departamento legal ha recibido para su cobranza "
        f"{cob.tipo} los documentos que se detallan a continuación, enviados "
        f"por nuestro cliente {razon}, Rut. :{rut_cliente}, Estudio: {estudio}."
    )
    _linea(doc, cuerpo, justificar=True)
    _linea(doc, "")

    # Detalle del documento (columnas monoespaciadas como la carta).
    _linea(doc, f"         {'FECHA VCTO':<19}{'TIP.DOC':<17}{'NRO.DOC':<17}{'MONTO':<15}C.P.")
    fecha_vcto = (cob.fecha_vencimiento_pagare.strftime("%d-%m-%Y")
                  if cob.fecha_vencimiento_pagare else "—")
    tip = ABREV_DOCUMENTO.get(cob.tipo_documento or "pagare", "OT")
    nro = cob.numero_pagare or (cob.id_clinica or "—")
    _linea(doc, f"         {fecha_vcto:<19}{tip:<17}{nro:<17}{_miles(nominal):<15}0")
    _linea(doc, "")

    # Cuadro resumen.
    _linea(doc, f"          {'NOMINAL':<13}{'GASTOS':<12}{'ABONOS':<12}{'HONORARIOS':<15}{'INTERESES':<14}SALDO")
    _linea(doc, f"          {_miles(nominal):<13}{_miles(gastos):<12}{_miles(abonos):<12}"
                f"{_miles(honorarios):<15}{_miles(intereses):<14}{_miles(saldo)}")
    _linea(doc, "")
    _linea(doc, "")

    # Formas de pago (bloque Redsalud tal cual; genérico para otros clientes).
    _linea(doc, f"    Formas de pago {estudio}:", tam=9, negrita=True)
    if "SALUD" in estudio:
        _linea(doc, "    - Directo en cajas en Clínica de su atención", tam=9)
        _linea(doc, "    - Transferencia electrónica a la cuenta de la clínica.", tam=9)
        _linea(doc, "    - Medios de pago electrónicos https://www.redsalud.cl/pagos-enlinea "
                    "(10 cuotas sin interés con tarjeta de crédito Banco Estado y hasta 10 "
                    "cuotas sin interés con tarjeta de crédito)", tam=9)
    else:
        _linea(doc, "    - Transferencia electrónica a la cuenta del cliente.", tam=9)
        _linea(doc, "    - Coordinar directamente con nuestro estudio al fono (32)2450990.", tam=9)
    _linea(doc, "    * una vez pagada su deuda debe enviarnos el comprobante por esta vía "
                "para registrarlo y terminar su cobranza", tam=9)
    _linea(doc, "")
    _linea(doc, "")
    _linea(doc, "    Sin otro particular le saluda atentamente,", tam=9)
    _linea(doc, "")
    _linea(doc, "ESTUDIO JURÍDICO HADAD & ASOCIADOS", tam=9, negrita=True,
           alineacion=WD_ALIGN_PARAGRAPH.RIGHT)

    # Pie de página con los datos de contacto del estudio (como la carta).
    pie = seccion.footer.paragraphs[0]
    r = pie.add_run(PIE_CARTA)
    r.font.name = "Courier New"
    r.font.size = Pt(8.5)

    nombre = f"estado_cuenta_cobranza_{cob.numero}_{date.today().isoformat()}.docx"
    return _responder_docx(doc, nombre)
